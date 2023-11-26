# Python标准库
import os
import shutil
import sqlite3
import subprocess
import tempfile
from collections import Counter
from io import BytesIO
from pathlib import Path
from re import split
from urllib import parse

# Django相关导入
from django.conf import settings
from django.contrib import messages
from django.core.files.storage import default_storage
from django.db.models import Q
from django.http import HttpResponse, FileResponse, Http404,JsonResponse
from django.shortcuts import render, redirect

# 第三方库
import docx
import jieba
import pandas as pd
import plotly.express as px
from plotly.offline import plot
from pyecharts.charts import WordCloud

# 本地应用/项目内部模块
from .forms import ArchiveModelForm  # 使用ModelForm组件来处理表单校验、数据保存和模板简写
from .utils.pagination import Pagination  # 自定义分页组件
from archives_app import models  # 导入模型用于数据库操作



def home(request, pages=1):
    # 初始化查询条件
    query = Q()

    # 搜索字段列表
    search_fields = ['num', 'name', 'line', 'station', 'theme', 'characteristic', 'effect']

    for field in search_fields:
        # 获取搜索参数
        value = request.GET.get(field, '')
        if value:
            # 动态构建查询条件
            query &= Q(**{f"{field}__icontains": value})

    # 根据查询条件获取数据并按照案例编号排序
    archives_queryset = models.Archives.objects.filter(query).order_by('num')
    # 利用分页组件对数据进行分页并传递给模板
    page_object = Pagination(request, archives_queryset, page_size=6, page_param="pages", plus=5)

    # 获取除id和fileName外的所有字段名
    exclude_columns = ['id', 'fileName']
    # 获取字段的verbose_name，并传递给模板
    columns = [f.verbose_name for f in models.Archives._meta.fields if f.name not in exclude_columns]

    context = {
        'columns': columns,
        'archives_queryset': page_object.page_queryset,
        'page_string': page_object.html(),
    }

    return render(request, 'home.html', context)

def archives_edit(request, edit_id):
    """编辑档案"""
    
    # 获取要编辑的数据id（Primary_key）
    row_obj = models.Archives.objects.filter(id=edit_id).first()
    
    # GET请求
    if request.method == "GET":
        # 创建表单实例并展示各字段的当前值
        form = ArchiveModelForm(instance=row_obj)
        return render(request, "archives_edit.html", {"form": form})
    
    # POST请求
    form = ArchiveModelForm(data=request.POST, instance=row_obj)
    if form.is_valid():
        form.save()
        return redirect("http://127.0.0.1:8000/")

    return render(request, "archives_edit.html", {"form": form}) 

def convert_to_docx(file_path, original_file_name):
    ''' 将 .doc 文件转换为 .docx 文件并返回新文件的路径'''
    # # 处理中文文件名的编码问题
    original_file_name_utf8 = original_file_name.encode('utf-8', errors='ignore').decode('utf-8')
    base_file_name = os.path.splitext(original_file_name_utf8)[0]
    docx_file_name = f"{base_file_name}.docx"
    docx_path = os.path.join(os.path.dirname(file_path), docx_file_name)

    command = [
      'S:\\LibreOfficePortable\\App\\libreoffice\\program\\soffice', # 若运行在其他人的电脑环境，需配置好LibreOfficePortable的环境路径
      '--headless',
      '--convert-to', 'docx',
      '--outdir', tempfile.gettempdir(),
      file_path
    ]

    try:
        # 运行命令
        result = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE)

        # 解码输出
        stdout_output = result.stdout.decode('utf-8', errors='ignore')
        stderr_output = result.stderr.decode('utf-8', errors='ignore')

        # 打印输出和错误，以便调试
        print("STDOUT:", stdout_output)
        print("STDERR:", stderr_output)

        # 检查命令是否成功执行
        if result.returncode != 0:
            print(f"文件 {original_file_name_utf8} 转换失败")
            return None

        # 如果转换成功，文件将保存在临时目录
        temp_docx_path = os.path.join(tempfile.gettempdir(), docx_file_name)

        # 将文件从临时目录移动到期望的位置
        shutil.move(temp_docx_path, docx_path)
        # 删除原始的 .doc 文件
        os.remove(file_path)  # 确保这是安全的操作，并且此时不再需要原始文件
        return docx_path

    except Exception as e:
        print(f"发生错误: {e}")
        return None

def process_word_document(file_path, file_name):
    """
    从 Word 文档中提取数据 最后返回一个列表
    其中包含从表格中提取的数据和文件名
    """
    try:
        doc = docx.Document(file_path)
        if not doc.tables:
            print(f"No tables found in {file_name}.")
            return None

        # 获取第一个表格的第二列的所有单元格
        table1 = doc.tables[0]
        # 获取第二列的所有单元格
        columns_cells = table1.columns[1].cells
        # 使用生成器表达式获取第二列的所有单元格的文本 比列表推导式更节省内存
        data = (cell.text for cell in columns_cells)
        # 将文件名添加到列表中
        data = [*data, file_name]
        return data
    except Exception as e:
        print(f"Error processing table in {file_name}: {e}")
        return None


def archives_generate(request):
    ''' 从上传的 .docx 文件中提取数据并保存到数据库 '''

    # Form表单提交过来
    if request.method == 'POST':
        # 确保 MEDIA_ROOT 目录存在
        media_root = settings.MEDIA_ROOT
        if not os.path.exists(media_root):
            os.makedirs(media_root)

        # request.FILES请求发来的是FILES文件, getlist得到文件夹下的批量文件file_list或批量上传的文件
        file_list = request.FILES.getlist("file")
        for upload_file in file_list:
            # 获取文件名后续保存进数据库中的字段中
            file_name = upload_file.name
            # 如果文件名以~$开头，则跳过（以~$开头的文件是Word文档的临时文件）
            if file_name.startswith('~$'):
                continue
            # 拼接文件路径
            file_path = os.path.join(media_root, file_name)

            # 保存文件到 MEDIA_ROOT
            # Looping over UploadedFile.chunks() instead of using read() ensures that large files don’t overwhelm  system’s memory.
            with open(file_path, 'wb+') as destination:
                # 从内存中分块读取文件并写入到磁盘
                for chunk in upload_file.chunks():
                    destination.write(chunk)

            if file_name.endswith(".doc"):
                # 调用函数转换 .doc 文件到 .docx
                docx_path = convert_to_docx(file_path, file_name)
                # 读取转变后的 .docx 文件中的表格
                if docx_path:
                    data = process_word_document(docx_path, os.path.basename(docx_path)) # 更新 .doc为 .docx文件名
                else:
                    print(f"Failed to convert {file_name}")
                    continue
            # 如果文件是 .docx 格式，直接读取表格
            elif file_name.endswith(".docx"):
                data = process_word_document(file_path, file_name)
            else:
                print(f"Skipped non-doc/docx file: {file_name}")
                continue

            if data is None:
                continue
            # 创建模型实例并批量保存到数据库
            # 遍历提取的数据                
            data_dict = {key: value for key, value in zip(['num', 'name', 'line', 'station', 'theme', 'characteristic', 'effect', 'fileName'], data)}
            # 创建表单实例
            form = ArchiveModelForm(data_dict)
            # 检查表单是否有效
            if form.is_valid():
                # 检查数据库中是否已存在相同的记录
                exists = models.Archives.objects.filter(
                    num=data[0],
                    name=data[1],
                    line=data[2],
                    station=data[3]
                ).exists()

                # 如果不存在重复数据，则保存
                if not exists:
                    form.save()
                    messages.success(request, f"'{file_name}' uploaded successfully.")
                else:
                    messages.warning(request, f"'{file_name}' already exists. Not saved.")
            else:
                # 在消息框架中记录表单错误
                messages.error(request, f"'{file_name}' upload failed. Incomplete or incorrect data.")
                # 记录详细的错误日志，如果需要可以写入日志文件
                print(f"Form errors for {file_name}: {form.errors}")

        return redirect('http://127.0.0.1:8000/')
    else:
        return HttpResponse('This view can only handle POST requests.', status=405)


def archives_delete(request, delete_id, pages=1):
    '''
    传递id对特定的数据进行删除
    '''
    models.Archives.objects.filter(id=delete_id).delete()
    return redirect('http://127.0.0.1:8000')


def delete_all(request):
    '''
    删除所有数据
    '''
    models.Archives.objects.all().delete()
    return redirect('http://127.0.0.1:8000')


def archives_download(request, id, pages=1):
    '''
    根据id索引到对应的文件files并下载
    '''
    download_obj = models.Archives.objects.filter(id=id).first()
    if not download_obj:
        raise Http404("Archive object does not exist.")

    # 获取文件名
    file_name = download_obj.fileName
    
    # 构建文件路径
    file_path = os.path.join(settings.MEDIA_ROOT, file_name)
    if not os.path.exists(file_path):
        raise Http404("File does not exist.")

    try:
        # 打开文件并创建响应
        file = default_storage.open(file_path, 'rb')
        response = FileResponse(file)
        response['content_type'] = "application/octet-stream"
        response['Content-Disposition'] = f'attachment; filename={parse.quote(file_name)}'
        return response
    except Exception as e:
        raise Http404(f"An error occurred while trying to download the file: {e}")
    
def archives_export(request):
    '''
    导出数据到Excel并提供下载且仅当Archives表存在时
    '''
    # 数据库路径获取
    db_path = settings.DATABASES['default']['NAME']
    # 连接到数据库
    conn = sqlite3.connect(db_path)

    # 检查Archives表是否存在
    cursor = conn.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='Archives';")
    if not cursor.fetchone():  # 如果没有找到Archives表
        conn.close()
        return JsonResponse({'error': 'Archives表不存在，无法导出数据。'}, status=404)

    # 读取表格
    df = pd.read_sql_query("SELECT * FROM Archives", conn)
    # 删除df中的第一列
    df.drop(df.columns[0], axis=1, inplace=True)

    # 关闭数据库连接
    conn.close()
    
    # 使用BytesIO作为Excel文件的缓冲区
    excel_file = BytesIO()

    # 将 DataFrame 保存到缓冲区中的 Excel 文件
    with pd.ExcelWriter(excel_file, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)

    # 重置文件指针到开始
    excel_file.seek(0)

    # 创建HTTP响应，设置内容类型为Excel文件
    response = HttpResponse(excel_file.read(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="Archives.xlsx"'

    return response

def load_stopwords(filenames):
    '''加载本地的停用词'''
    stopwords = set()
    # 静态文件目录路径
    static_dir = os.path.join(settings.BASE_DIR, 'static')
    for filename in filenames:
        # 构建文件的完整路径
        file_path = Path(static_dir) / 'plugins' / 'stopwords-master' / filename
        with open(file_path, 'r', encoding='utf-8') as file:
            stopwords.update([line.strip() for line in file.readlines()])
    return stopwords

# 停用词文件名列表
stopwords_files = [
    'baidu_stopwords.txt',
    'cn_stopwords.txt',
    'hit_stopwords.txt',
    'scu_stopwords.txt'
]

# 加载停用词
STOPWORDS = load_stopwords(stopwords_files)

def create_wordcloud_pyecharts(data):
    '''使用pyecharts创建词云图'''
    
    # 将所有文本合并为一个长字符串
    text = ' '.join(data)

    # 使用jieba进行中文分词
    words = jieba.cut(text)
    filtered_words = [word for word in words if word not in STOPWORDS]

    # 统计词频
    counts = Counter(filtered_words)

    # 准备词云图数据
    wordcloud_data = [(word, freq) for word, freq in counts.items()]

    # 创建词云图
    wordcloud = WordCloud()
    wordcloud.add("", wordcloud_data, word_size_range=[20, 100])
    return wordcloud.render_embed()

# 定义一个辅助函数来创建饼图
def create_pie_chart(data, title):
    counts = pd.Series(data).value_counts().reset_index()
    counts.columns = ['category', 'count']
    fig = px.pie(counts, values='count', names='category', title=title)
    fig.update_layout(
        title={
            'text': title,
            'y':0.9,
            'x':0.5,
            'xanchor': 'center',
            'yanchor': 'top'
        },
        title_font=dict(size=24),
    )
    fig.update_traces(
        hoverinfo='label+percent+name',
        textinfo='none',
        hovertemplate='类别: %{label}<br>计数: %{value}<extra></extra>'
    )
    return plot(fig, output_type='div', include_plotlyjs=False)

def archives_visualization(request):
    '''
    数据可视化
    '''
    # 从模型中获取字段数据
    themes = models.Archives.objects.values_list('theme', flat=True)
    lines_data = models.Archives.objects.values_list('line', flat=True)
    stations_data = models.Archives.objects.values_list('station', flat=True)
    characteristic_data = models.Archives.objects.values_list('characteristic', flat=True)

    # 处理线路数据
    lines = []
    filters = ['全网车站', '全网', '全线车辆', '全线车站', '全网车辆']
    for line in lines_data:
        if line in filters:
            lines.append('全网')
        else:
            # 分割线路，考虑“、”和“/”作为分隔符
            separated_lines = split(r'[、/]', line)
            for l in separated_lines:
                if l:
                    # 如果线路名称不以“号线”结尾，则添加“号线”后缀
                    l = l if l.endswith('号线') else l + '号线'
                    lines.append(l)

    # 处理车站数据，逻辑与线路类似
    stations = []
    for station in stations_data:
        if station in filters:
            stations.append('全网')
        else:
            stations.append(station)
    wordcloud_html = create_wordcloud_pyecharts(characteristic_data)

    # 创建饼图并准备传递给模板的上下文
    context = {
        'theme_pie_chart_div': create_pie_chart(themes, '主题分布'),
        'line_pie_chart_div': create_pie_chart(lines, '线路分布'),
        'station_pie_chart_div': create_pie_chart(stations, '车站分布'),
        'wordcloud_html': wordcloud_html
    }

    # 将上下文传递到模板
    return render(request, 'visualization.html', context)
